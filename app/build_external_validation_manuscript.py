from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REFERENCE_ADDITIONS = [
    "26. Chivukula N, Vashishth S, Kandasamy P, Madgaonkar SR, Samal A (2026) DEDuCT 3.0: an enhanced and expanded FAIR-compliant resource and toxicology knowledge graph for endocrine disrupting chemicals. bioRxiv. doi:10.64898/2026.01.23.701267.",
    "27. Health Canada Pest Management Regulatory Agency. Pesticide Product Information Database open data extracts. https://pest-control.canada.ca/pesticide-registry-api/api/extract/. Accessed 6 Aug 2026.",
    "28. European Commission. Cosmetic ingredient database (CosIng). https://single-market-economy.ec.europa.eu/sectors/cosmetics/cosmetic-ingredient-database_en. Accessed 16 Aug 2026.",
    "29. U.S. Food and Drug Administration. Inventory of Effective Food Contact Substance Notifications. https://www.fda.gov/food/packaging-food-contact-substances-fcs/inventory-effective-food-contact-substance-fcs-notifications. Accessed 16 Aug 2026.",
]

DISPLAY_NAMES = {
    "endocrine_disruptors": "Endocrine disruptors",
    "flavor_fragrance": "Flavor and fragrance",
    "pesticides": "Pesticides",
    "surfactants": "Surfactants",
}

SOURCE_CITATIONS = {
    "DEDuCT v3 I-III": 26,
    "Health Canada PMRA PPID": 27,
}

ROOT = Path(__file__).resolve().parent.parent
MODEL_IDS = {
    **{category: f"final_{category}" for category in DISPLAY_NAMES},
    "endocrine_disruptors": "han_endocrine_disruptors",
}


def read_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def load_current_models() -> dict[str, dict[str, Any]]:
    model_dir = ROOT / "app" / "data" / "models"
    return {
        category: json.loads((model_dir / f"{MODEL_IDS[category]}.json").read_text(encoding="utf-8"))
        for category in DISPLAY_NAMES
    }


def selected_descriptors(config: dict[str, Any]) -> list[str]:
    return list(config.get("selected_props", config.get("descriptor_ranges", {})))


def selected_patterns(config: dict[str, Any]) -> list[str]:
    return list(config.get("selected_patterns", config.get("smarts_patterns", {})))


def model_weight_text(config: dict[str, Any]) -> str:
    if config.get("model_type") == "han_edc":
        weights = config.get("weights", {})
        return (
            f"scaffold={float(weights.get('scaffold_score', 0)):.3f};"
            f"similarity={float(weights.get('similarity_score', 0)):.3f};"
            f"smarts={float(weights.get('smarts_score', 0)):.3f}"
        )
    return f"{float(config.get('best_w', 0.5)):.4g}"


def benchmark_maps(
    ablation_rows: list[dict[str, str]], bootstrap_rows: list[dict[str, str]]
) -> tuple[dict[str, dict[str, str]], dict[str, dict[str, str]]]:
    combined = {
        row["category"]: row for row in ablation_rows if row["component"] == "Combined score"
    }
    bootstrap = {
        row["category"]: row for row in bootstrap_rows if row["metric"] == "qed_auc"
    }
    return combined, bootstrap


def refresh_markdown_core_tables(
    text: str,
    ablation_rows: list[dict[str, str]],
    bootstrap_rows: list[dict[str, str]],
) -> str:
    combined, qed = benchmark_maps(ablation_rows, bootstrap_rows)
    models = load_current_models()
    lines = text.splitlines()
    table_rows: dict[str, tuple[int, int]] = {}
    for index, line in enumerate(lines):
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if not cells or cells[0] not in DISPLAY_NAMES.values():
            continue
        category = next(key for key, display in DISPLAY_NAMES.items() if display == cells[0])
        if len(cells) == 8 and category in combined:
            score_auc = float(combined[category]["auc"])
            balanced = float(combined[category]["maximum_balanced_accuracy"])
            qed_auc = float(qed[category]["estimate"])
            cells[4:] = [f"{score_auc:.4f}", f"{balanced:.4f}", f"{qed_auc:.4f}", f"{score_auc-qed_auc:.4f}"]
            lines[index] = "| " + " | ".join(cells) + " |"
            table_rows.setdefault("benchmark", (index, len(cells)))
        elif len(cells) == 4:
            config = models[category]
            cells[1:] = [
                ", ".join(selected_descriptors(config)) or "none",
                ", ".join(selected_patterns(config)) or "none",
                model_weight_text(config),
            ]
            lines[index] = "| " + " | ".join(cells) + " |"
            table_rows.setdefault("composition", (index, len(cells)))
    for table_type, columns in (("benchmark", 8), ("composition", 4)):
        if table_type not in table_rows or "flavor_fragrance" not in combined:
            continue
        anchor, _ = table_rows[table_type]
        display = DISPLAY_NAMES["flavor_fragrance"]
        if any(line.startswith(f"| {display} ") for line in lines):
            continue
        if columns == 8:
            score_auc = float(combined["flavor_fragrance"]["auc"])
            balanced = float(combined["flavor_fragrance"]["maximum_balanced_accuracy"])
            qed_auc = float(qed["flavor_fragrance"]["estimate"])
            values = [display, "3531", "12679", "12679", f"{score_auc:.4f}", f"{balanced:.4f}", f"{qed_auc:.4f}", f"{score_auc-qed_auc:.4f}"]
        else:
            config = models["flavor_fragrance"]
            values = [display, ", ".join(selected_descriptors(config)) or "none", ", ".join(selected_patterns(config)) or "none", model_weight_text(config)]
        insertion = anchor
        while insertion + 1 < len(lines) and lines[insertion + 1].startswith("|"):
            insertion += 1
        lines.insert(insertion + 1, "| " + " | ".join(values) + " |")
    return "\n".join(lines) + ("\n" if text.endswith("\n") else "")


def refresh_docx_core_tables(
    document: Document,
    ablation_rows: list[dict[str, str]],
    bootstrap_rows: list[dict[str, str]],
) -> None:
    combined, qed = benchmark_maps(ablation_rows, bootstrap_rows)
    models = load_current_models()
    for table in document.tables:
        for row in table.rows:
            category_name = row.cells[0].text.strip()
            if category_name not in DISPLAY_NAMES.values():
                continue
            category = next(key for key, display in DISPLAY_NAMES.items() if display == category_name)
            if len(row.cells) == 8 and category in combined:
                score_auc = float(combined[category]["auc"])
                balanced = float(combined[category]["maximum_balanced_accuracy"])
                qed_auc = float(qed[category]["estimate"])
                for cell, value in zip(
                    row.cells[4:],
                    (f"{score_auc:.4f}", f"{balanced:.4f}", f"{qed_auc:.4f}", f"{score_auc-qed_auc:.4f}"),
                ):
                    set_cell_text(cell, value)
            elif len(row.cells) == 4:
                config = models[category]
                values = (
                    ", ".join(selected_descriptors(config)) or "none",
                    ", ".join(selected_patterns(config)) or "none",
                    model_weight_text(config),
                )
                for cell, value in zip(row.cells[1:], values):
                    set_cell_text(cell, value)
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if "flavor_fragrance" not in combined or any(row.cells[0].text.strip() == DISPLAY_NAMES["flavor_fragrance"] for row in table.rows):
            continue
        if len(headers) == 8 and "AUC" in headers:
            score_auc = float(combined["flavor_fragrance"]["auc"])
            balanced = float(combined["flavor_fragrance"]["maximum_balanced_accuracy"])
            qed_auc = float(qed["flavor_fragrance"]["estimate"])
            values = [DISPLAY_NAMES["flavor_fragrance"], "3531", "12679", "12679", f"{score_auc:.4f}", f"{balanced:.4f}", f"{qed_auc:.4f}", f"{score_auc-qed_auc:.4f}"]
        elif len(headers) == 4:
            config = models["flavor_fragrance"]
            values = [DISPLAY_NAMES["flavor_fragrance"], ", ".join(selected_descriptors(config)) or "none", ", ".join(selected_patterns(config)) or "none", model_weight_text(config)]
        else:
            continue
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)


def replace_docx_inline_images(docx_path: Path, image_paths: list[Path]) -> None:
    document = Document(docx_path)
    if len(document.inline_shapes) < len(image_paths):
        raise ValueError("DOCX has fewer inline images than the requested replacements")
    replacements: dict[str, bytes] = {}
    for shape, image_path in zip(document.inline_shapes, image_paths):
        relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
        part = document.part.related_parts[relationship_id]
        replacements[str(part.partname).lstrip("/")] = image_path.read_bytes()
    temporary = docx_path.with_suffix(".images.tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as destination:
        for item in source.infolist():
            destination.writestr(item, replacements.get(item.filename, source.read(item.filename)))
    temporary.replace(docx_path)


def integer(row: dict[str, str], key: str) -> int:
    value = row.get(key, "")
    return int(float(value)) if value else 0


def percent(row: dict[str, str]) -> str:
    value = row.get("positive_recovery_fraction", "")
    return "not estimable" if not value else f"{100 * float(value):.1f}%"


def evidence_rows(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    return [
        row
        for row in rows
        if row.get("category") in DISPLAY_NAMES
        and integer(row, "true_external_candidates") > 0
        and integer(row, "scored_true_external") > 0
    ]


def primary_rows(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    return [row for row in evidence_rows(rows) if row.get("evidence_scope") == "external positive comparison"]


def source_list(rows: list[dict[str, str]]) -> str:
    sources = [
        f"{row['external_source']} [{SOURCE_CITATIONS[row['external_source']]}]"
        for row in rows
        if row.get("category") in DISPLAY_NAMES
    ]
    if len(sources) == 1:
        return sources[0]
    if len(sources) == 2:
        return " and ".join(sources)
    return ", ".join(sources[:-1]) + f", and {sources[-1]}"


def prose_list(items: list[str]) -> str:
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return " and ".join(items)
    return ", ".join(items[:-1]) + f", and {items[-1]}"


def result_summary(rows: list[dict[str, str]]) -> str:
    primary = primary_rows(rows)
    statements = []
    for row in primary:
        statements.append(
            f"{DISPLAY_NAMES[row['category']]}: {integer(row, 'recovered_above_threshold')}/"
            f"{integer(row, 'scored_true_external')} ({percent(row)})"
        )
    return "; ".join(statements)


def overlap_summary(rows: list[dict[str, str]]) -> str:
    statements = []
    for row in rows:
        shared = integer(row, "shared_unique_parent_structures")
        external = integer(row, "external_unique_parent_structures")
        coverage = 100 * float(row["external_parent_coverage_fraction"])
        statements.append(
            f"{row['category_display']}: {shared:,}/{external:,} ({coverage:.1f}%)"
        )
    return "; ".join(statements)


def overlap_evidence_text(rows: list[dict[str, str]]) -> str:
    return (
        "At the level of unique standardized parent structures, direct comparison of the external sets with the "
        f"corresponding PubChem-derived category sets gave the following overlap: {overlap_summary(rows)} (Figure 5). "
        "Thus, the PubChem-derived sets contained most of the evaluated DEDuCT endocrine-disruptor structures and "
        "structure-resolved PMRA pesticide substances. These results "
        "describe overlap among curated molecular datasets, not coverage of marketed products. In particular, the PMRA "
        "percentage applies only to records for which a structure was successfully resolved."
    )


def apply_terminology(text: str) -> str:
    replacements = (
        ("PubChem category membership", "the compounds assigned to each PubChem category"),
        ("PubChem-derived category memberships", "PubChem-derived category sets"),
        ("PubChem-derived memberships", "PubChem-derived category sets"),
        ("PubChem membership list", "PubChem-derived category set"),
        ("external positive memberships", "external positive category assignments"),
        ("positive memberships", "positive category assignments"),
        ("independently curated memberships", "independently curated category assignments"),
        ("database memberships", "database category assignments"),
        ("Database Memberships", "Database Category Assignments"),
        ("membership in an unrelated category", "assignment to an unrelated category"),
        ("derived from PubChem membership", "defined from compounds assigned to PubChem categories"),
    )
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def methods_text(rows: list[dict[str, str]]) -> tuple[str, str]:
    return (
        "The final score definitions and decision thresholds were fixed before the external database comparison. "
        f"Positive category assignments were obtained from {source_list(rows)}. The comparison applied each score only "
        "to molecules assigned to the corresponding category; it did not treat assignment to an unrelated category as "
        "a positive result. This analysis examines transfer of the fixed rules beyond the PubChem-derived category sets "
        "and is not presented as external validation of a fitted QSAR or machine-learning model.",
        "External structures were standardized with RDKit cleanup, fragment-parent selection, and charge normalization. "
        "An external record was excluded when either its standardized whole structure or the connectivity block of its "
        "parent InChIKey matched any molecule used in any final positive or negative-source construction set. Repeated "
        "external parent structures were also removed after the first occurrence. Only successfully resolved, nonoverlapping "
        "structures were scored, without parameter or threshold refitting. Direct overlap between each external source "
        "and the corresponding PubChem-derived target set was also counted at the level of unique standardized parent "
        "structures; the denominator for the reported representation percentage was the number of unique external parent "
        "structures. Because the external sources provide positive category assignments but not verified category negatives, "
        "the reported score-performance measures are the number and fraction of external positives at or above the "
        "pre-established threshold and the score distribution; ROC AUC, specificity, and balanced accuracy cannot be "
        "estimated from these positive-only data.",
    )


def result_paragraphs(rows: list[dict[str, str]]) -> tuple[str, str]:
    main = primary_rows(rows)
    first = (
        f"After structure resolution, construction-set overlap removal, and external duplicate removal, the fixed scores "
        f"were applied to {sum(integer(row, 'scored_true_external') for row in main):,} nonoverlapping external positives "
        f"from {len(main)} category-source pairs (Table 3). Recovery at the pre-established thresholds was {result_summary(rows)}. "
        "These values measure agreement with external positive category assignments after molecule-level separation; they are not "
        "classification accuracies because the external datasets do not supply verified negatives."
    )
    source_overlap = [
        f"{DISPLAY_NAMES[row['category']]} ({row['external_source']})"
        for row in rows
        if row.get("category") in DISPLAY_NAMES
        and row.get("evidence_scope") != "external positive comparison"
    ]
    not_evaluable = [
        f"{DISPLAY_NAMES[row['category']]} ({row['external_source']})"
        for row in rows
        if row.get("category") in DISPLAY_NAMES and integer(row, "scored_true_external") == 0
    ]
    second = (
        "The external comparison was category dependent. The lower recovery "
        "for DEDuCT endocrine disruptors shows that the PubChem-derived endocrine score does not cover much of the broader "
        "literature-defined endocrine-disruption space. Only 15 nonoverlapping pesticide structures could be scored after "
        "incomplete identifier resolution and extensive overlap removal, and three reached the threshold; this small, selected "
        "subset cannot establish pesticide transfer. These differences are scientifically informative and prevent a universal "
        "robustness claim."
    )
    if source_overlap:
        second += " " + prose_list(source_overlap) + " were treated only as source-lineage consistency checks, not as independent label validation."
    if not_evaluable:
        second += " No recovery estimate was made for " + prose_list(not_evaluable) + " because no scored nonoverlapping structures remained."
    return first, second


def discussion_text(rows: list[dict[str, str]], overlap_rows: list[dict[str, str]]) -> list[str]:
    _, external_interpretation = result_paragraphs(rows)
    return [
        "The main contribution of this study is bounded but defensible. The present evidence does not establish category-likeness scoring as a wholly new concept, because QED, synthetic-accessibility scoring, natural-product-likeness, ChEMBL-likeness, structural-alert methods, and pesticide-focused QEP already show that explicit scoring can be chemically useful in defined domains [1-9]. The contribution supported here is narrower: one scoring strategy based on molecular descriptors and structural patterns can be applied across a mixed panel of broad PubChem product classes, benchmarked against the same comparison background, and compared with external positive category assignments after structure-level overlap removal.",
        "The most critical methodological choice is the negative set, not the existence of named motifs by itself. Descriptor windows and pattern enrichment can appear highly discriminating when the comparison background is chemically remote. The restrained negative-selection procedure used here keeps neighboring-category chemistry in the background and excludes exact target overlap, duplicate background structures, and near-positive analogues identified by Morgan-fingerprint similarity. The scores should therefore be interpreted together with Figure 1 and Table 1, which document the construction of the comparison sets.",
        "The interpretability claim is narrower than a mechanistic explanation claim. The scores do not establish why a molecule has a regulatory or commercial use. They expose the information used by each scoring function: descriptor intervals, named SMARTS motifs, and the relative contributions of those terms. Table 2 and Figure 4 show this information directly and identify structural patterns shared by, or enriched in, particular categories. The functions are therefore inspectable rule-based scores rather than black-box classifiers.",
        "The retained benchmark profile comprises surfactants, endocrine disruptors, pesticides, and a merged flavor-and-fragrance function. Pesticides exceeded QED and the reimplemented QEP comparators, although pesticide-likeness has the strongest direct precedent and the novelty claim must remain modest. Cosmetics, food-contact substances, food additives, and solvents were excluded after exact-overlap-controlled cross-category revalidation showed inadequate threshold specificity.",
        external_interpretation,
        overlap_evidence_text(overlap_rows),
        "The external positive-set analysis is useful even though no statistical learning algorithm was fitted. It does not serve the usual purpose of estimating generalization error for a QSAR model. Instead, fixing the scoring rules and thresholds before applying them to independently curated category assignments tests whether a PubChem-derived category definition transfers to another annotation source. Molecule-level overlap removal is essential for that question because an apparently external database may contain the same structures as the construction data.",
        "The QED study provides a relevant precedent for interpreting such evidence. Its 771-compound DrugBank benchmark contained 554 structures identical to the drugs used to derive QED, with another 30 compounds exceeding a Tanimoto similarity of 0.8. The authors disclosed this dependence, removed 475 highly similar compounds from the PDB-derived negative set, retained 10,250 comparison compounds, and assessed relative performance using receiver operating characteristic curves, sensitivity, specificity, and the Matthews correlation coefficient [1]. Accordingly, the present source overlaps are reported explicitly and the nonoverlapping subsets are treated as relative transfer tests rather than fully independent validation sets.",
        "Several limitations remain. The main benchmark uses constructed overlapping comparison chemistry rather than universally verified true negatives. The external sources contain positive category assignments only, so they cannot support external ROC AUC, specificity, or balanced accuracy. Annotation-lineage independence cannot be proven merely from different database names, and source-overlapping datasets were therefore separated from independent positive comparisons. Structure resolution was incomplete for some identifier lists, which can introduce selection bias. The score library is constrained by predefined SMARTS motifs, and recurring broad motifs are not exclusive category identifiers. Future work should extend external positive comparisons to additional independently curated sources and compare the fixed motif library with Murcko-derived, BRICS-derived, or hybrid fragment candidates under the same matched design [20-23].",
    ]


def conclusions_text(rows: list[dict[str, str]]) -> list[str]:
    main = primary_rows(rows)
    return [
        "Scoring-function development was attempted for all eleven original PubChem categories. Rigorous cross-category screening retained four reportable functions: endocrine disruptors, pesticides, surfactants, and a merged flavor-and-fragrance function. Animal-drug and human-drug functions were excluded because QED already represents general drug-likeness; cosmetics, food-contact-substance, food-additive, and solvent functions were excluded for inadequate threshold specificity. Flavoring agents and fragrances were merged because their positive sets shared 1,071 structures and the separate scores used near-identical descriptor and motif logic.",
        f"When the fixed rules were applied without refitting to {sum(integer(row, 'scored_true_external') for row in main):,} nonoverlapping external positives, recovery was strongly category dependent ({result_summary(rows)}). The endocrine-disruptor result defines an important limit on coverage, and the small pesticide set limits inference for that category. The defensible conclusion is therefore not universal external validity, but that explicit descriptor and structural-pattern scores can provide reusable, inspectable category evidence whose transfer must be established separately for each category and annotation source.",
    ]


def abstract_text(rows: list[dict[str, str]]) -> str:
    main = primary_rows(rows)
    return (
        "Accessible continuous scores remain uncommon for broad chemical product categories. We attempted scoring-function development for eleven categories defined from compounds assigned to PubChem categories and subjected every candidate to exact-overlap-controlled cross-category screening. Four functions survived: endocrine disruptors, pesticides, surfactants, and a merged flavor-and-fragrance function. The merged target was justified by 1,071 shared structures, representing 46% of each original positive set, and showed held-out AUCs of 0.842–0.868 across three hash folds. Animal-drug and human-drug functions were excluded because QED already represents general drug-likeness; cosmetics, food-contact-substance, food-additive, and solvent functions were excluded for inadequate threshold specificity. For each retained target, the comparison background was filtered for target overlap, duplicate structures, and near-positive Morgan-fingerprint neighbors. The fixed rules were additionally applied, without refitting, to "
        f"{sum(integer(row, 'scored_true_external') for row in main):,} external positive structures remaining after overlap control. "
        f"Recovery at the pre-established thresholds was category dependent ({result_summary(rows)}). Each function can be decomposed into explicit descriptor intervals and named structural patterns. These findings support the use of inspectable category scores while showing that transfer to external annotation sources must be assessed separately for each chemical category."
    )


def markdown_table(rows: list[dict[str, str]]) -> str:
    lines = [
        "Table 3. Comparison with external positive category assignments after structure resolution and overlap control. Recovery is the fraction of scored external positives at or above the pre-established category threshold.",
        "",
        "| Category | External source | Initial records | Resolved structures | Construction overlap excluded | External duplicates excluded | Scored nonoverlapping positives | Recovered above threshold | Recovery | Evidence scope |",
        "| --- | --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: | --- |",
    ]
    for row in rows:
        if row.get("category") not in DISPLAY_NAMES:
            continue
        lines.append(
            f"| {DISPLAY_NAMES[row['category']]} | {row['external_source']} | {integer(row, 'raw_candidates')} | "
            f"{integer(row, 'resolved_structures')} | {integer(row, 'overlap_excluded')} | "
            f"{integer(row, 'duplicate_external_excluded')} | {integer(row, 'scored_true_external')} | "
            f"{integer(row, 'recovered_above_threshold')} | {percent(row)} | {row['evidence_scope']} |"
        )
    return "\n".join(lines)


def analysis_methods_text() -> str:
    return (
        "Four post hoc analyses were performed with the released score definitions and decision thresholds held fixed; "
        "no score was refitted. First, each of the four retained scores was applied to every retained positive category set, "
        "and the fraction at or above the corresponding frozen threshold was calculated. This threshold-based endpoint "
        "was used because the independently calibrated score scales are not directly comparable by their raw medians. "
        "Because the thresholds also have different false-positive rates, a second matrix subtracted each score's "
        "threshold-positive fraction in its retained comparison set from the fraction observed in each source positive set. Second, "
        "the descriptor, structural-evidence, and combined components were evaluated by ROC AUC against the same "
        "retained comparison structures. For the endocrine-disruptor score, the structural-evidence component comprised "
        "the scaffold, fingerprint-similarity, and SMARTS terms. Third, uncertainty was estimated by independently "
        "resampling positive and retained comparison structures in 1,000 stratified bootstrap replicates. Percentile "
        "95% confidence intervals were calculated for score AUC, QED AUC, their paired difference, balanced accuracy, "
        "and Matthews correlation coefficient at the frozen threshold. Finally, score distributions were compared "
        "among the PubChem-derived category sets, retained comparison structures, and nonoverlapping external positive "
        "sets for endocrine disruptors and pesticides. Separately, exact canonical-structure overlap was calculated for "
        "all 55 pairs among the eleven original positive sets. Pairwise shared counts, directional coverage fractions, "
        "Jaccard similarity, and the number of category assignments per unique structure were reported in Supporting "
        "Information Tables S2–S3 and Figures S2–S3."
        "\n\nAfter the cross-category analysis identified weak specificity, the animal-drug, human-drug, "
        "food-contact-substance, cosmetics, and food-additive functions were rebuilt sequentially. For each target, "
        "canonical target overlaps were removed from the other ten category sets, but near-positive structures were "
        "not removed by Tanimoto similarity. These chemically difficult cross-category structures were divided by a "
        "deterministic molecular hash into training and held-out sets. Candidate functions retained the same explicit "
        "descriptor-range and structural-pattern architecture. Promotion required improvement on held-out cross-category "
        "separation without a material loss on the original retained-comparison benchmark. The only promising target, "
        "cosmetics, was confirmed across three held-out hash folds and then rebuilt from all available structures. "
        "Animal-drug and human-drug functions were subsequently excluded from the retained panel because QED already "
        "represents generic drug-likeness and those broad categories did not provide sufficiently distinct product-use evidence. "
        "For final publication screening, exact canonical overlaps with each score's target set were removed separately from "
        "every other source category. A score was retained only when its threshold-positive response on its own positive set "
        "exceeded its response on every exact-overlap-excluded source set and its AUC against the pooled exact-overlap-excluded "
        "cross-category background exceeded 0.5."
    )


def sequential_rebuild_results() -> str:
    return (
        "Exact-structure overlap was pervasive across the original category panel. Of 19,934 unique canonical structures, "
        "4,555 (22.9%) occurred in at least two categories and 1,173 occurred in three categories. The strongest cluster "
        "comprised flavoring agents, food additives, and fragrances: flavoring agents and food additives shared 1,752 "
        "structures (75.6% and 62.4% of the respective sets), food additives and fragrances shared 1,113 (39.6% and "
        "48.7%), and flavoring agents and fragrances shared 1,071 (46.2% and 46.9%). A three-way intersection contained "
        "943 structures. Complete pairwise and multiplicity results are reported in Supporting Information Tables S2–S3 "
        "and Figures S2–S3. "
        "Exact-overlap-controlled revalidation showed that industrial co-use could not explain the observed cross-response. "
        "After structures assigned to cosmetics were removed from the surfactant set, the promoted cosmetics function still "
        "called 81.0% of the remaining surfactants positive, compared with 43.5% of cosmetics positives. The food-contact "
        "function called 100% of surfactants not assigned to food contact positive, compared with 81.9% of its own positives. "
        "The food-additive function likewise responded more often to flavoring-agent structures than to its own positive set. "
        "The solvent function called 71.8% of flavoring agents and 68.8% of fragrances positive after exact solvent overlap "
        "was removed, reflecting broad low-polarity aliphatic descriptor ranges rather than solvent-specific evidence. "
        "These four functions were therefore excluded rather than explained post hoc by plausible industrial overlap. "
        "Flavoring agents and fragrances followed a different path: their positive sets shared 1,071 structures (46% of "
        "each set), and their separate functions used the same descriptor family and aldehyde, cinnamate, and ester motifs. "
        "They were merged into a broader flavor-and-fragrance target. Across three held-out hash folds, the merged function "
        "gave AUCs of 0.842–0.868 and positive responses of 70.0–79.0%; its highest source-specific response was observed "
        "for exact-overlap-excluded food additives (31.7–36.6%)."
    )


def cross_category_results(rows: list[dict[str, str]]) -> str:
    diagonal = [row for row in rows if row["source_category"] == row["score_category"]]
    rates = [100 * float(row["fraction_at_or_above_threshold"]) for row in diagonal]
    diagonal_wins = 0
    adjusted_diagonal_wins = 0
    for diagonal_row in diagonal:
        source = diagonal_row["source_category"]
        maximum = max(
            float(row["fraction_at_or_above_threshold"])
            for row in rows
            if row["source_category"] == source
        )
        if abs(float(diagonal_row["fraction_at_or_above_threshold"]) - maximum) < 1e-12:
            diagonal_wins += 1
        adjusted_maximum = max(
            float(row["excess_fraction_over_retained_comparison"])
            for row in rows
            if row["source_category"] == source
        )
        if abs(float(diagonal_row["excess_fraction_over_retained_comparison"]) - adjusted_maximum) < 1e-12:
            adjusted_diagonal_wins += 1
    return (
        f"Across the complete positive sets, the fraction exceeding the intended score threshold ranged from "
        f"{min(rates):.1f}% to {max(rates):.1f}%. The intended score produced the highest threshold-positive "
        f"fraction for {diagonal_wins} of the {len(diagonal)} source sets (Figure 6A). Some non-intended scores "
        "therefore covered a source positive set more extensively than its intended score. This comparison is affected "
        "by unequal operating points. After subtracting each score's corresponding "
        f"retained-comparison response, the intended score was highest for {adjusted_diagonal_wins} of "
        f"{len(diagonal)} source sets (Figure 6B), although substantial off-diagonal response remained. These responses "
        "are consistent with both shared chemical features and nonexclusive product-use categories. "
        "The functions therefore provide parallel measures of category-associated structural evidence rather than "
        "a mutually exclusive, winner-take-all classification."
    )


def ablation_results(rows: list[dict[str, str]]) -> str:
    by_category: dict[str, dict[str, float]] = {}
    for row in rows:
        by_category.setdefault(row["category"], {})[row["component"]] = float(row["auc"])
    combined_wins = sum(
        values["Combined score"] >= max(values["Descriptor component"], values["Structural evidence component"])
        for values in by_category.values()
    )
    endocrine = by_category["endocrine_disruptors"]
    return (
        f"The combined score had the highest AUC in {combined_wins} of {len(by_category)} categories (Figure 7). "
        f"For endocrine disruptors, structural evidence alone gave an AUC of "
        f"{endocrine['Structural evidence component']:.4f}, close to the combined-score AUC of "
        f"{endocrine['Combined score']:.4f}, indicating that explicit structural evidence accounted for most of the "
        "separation in this category."
    )


def bootstrap_results(rows: list[dict[str, str]]) -> str:
    deltas = [row for row in rows if row["metric"] == "auc_delta_vs_qed"]
    weakest = min(deltas, key=lambda row: float(row["estimate"]))
    return (
        f"All {len(deltas)} score-minus-QED AUC differences had bootstrap 95% confidence intervals above zero "
        f"(Table 4). The smallest estimated advantage was observed for {DISPLAY_NAMES[weakest['category']].lower()} "
        f"(ΔAUC {float(weakest['estimate']):.4f}, 95% CI {float(weakest['ci_lower_95']):.4f}–"
        f"{float(weakest['ci_upper_95']):.4f})."
    )


def external_distribution_results(
    rows: list[dict[str, str]], summary_rows: list[dict[str, str]]
) -> str:
    statements = []
    summary_by_category = {row["category"]: row for row in primary_rows(summary_rows)}
    for category in ("endocrine_disruptors", "pesticides"):
        values = sorted(float(row["score"]) for row in rows if row["category"] == category)
        midpoint = len(values) // 2
        median = values[midpoint] if len(values) % 2 else (values[midpoint - 1] + values[midpoint]) / 2
        summary = summary_by_category[category]
        statements.append(
            f"{DISPLAY_NAMES[category]}: median {median:.3f}, recovery {percent(summary)}, n = {len(values):,}"
        )
    return (
        "The nonoverlapping external positive sets showed category-dependent score distributions "
        f"({'; '.join(statements)}; Figure 8). The pesticide result is necessarily tentative because only 15 "
        "nonoverlapping resolved structures remained."
    )


def confidence_interval_markdown_table(rows: list[dict[str, str]]) -> str:
    metrics: dict[str, dict[str, dict[str, str]]] = {}
    for row in rows:
        metrics.setdefault(row["category"], {})[row["metric"]] = row
    lines = [
        "Table 4. Bootstrap estimates and 95% confidence intervals from 1,000 stratified resamples of positive and retained comparison structures.",
        "",
        "| Category | Score AUC (95% CI) | QED AUC (95% CI) | ΔAUC (95% CI) |",
        "| --- | ---: | ---: | ---: |",
    ]
    for category in DISPLAY_NAMES:
        category_metrics = metrics[category]
        values = []
        for metric in ("auc", "qed_auc", "auc_delta_vs_qed"):
            row = category_metrics[metric]
            values.append(
                f"{float(row['estimate']):.3f} ({float(row['ci_lower_95']):.3f}–{float(row['ci_upper_95']):.3f})"
            )
        lines.append(f"| {DISPLAY_NAMES[category]} | {' | '.join(values)} |")
    return "\n".join(lines)


def additional_discussion_text() -> list[str]:
    return [
        (
            "The full overlap matrix shows that the original source labels form an overlapping annotation system rather "
            "than eleven independent chemical classes. Nearly one quarter of unique structures had multiple assignments. "
            "The flavoring-agent, fragrance, and food-additive triangle was dominant, including 943 structures assigned "
            "to all three categories. This result supports merging flavoring agents and fragrances because their use and "
            "structural definitions were closely aligned, while also explaining why the broader food-additive function "
            "was nonspecific. Food additives were not merged because their scope includes nonsensory technological roles "
            "not represented by the retained flavor-and-fragrance motifs. The directional matrix also prevents asymmetric "
            "overlaps involving small sets, such as surfactants or solvents, from being mistaken for equivalence between "
            "the full categories (Supporting Information Figure S2)."
        ),
        (
            "The cross-category matrix places an important boundary on interpretation. Strong off-diagonal responses "
            "are chemically plausible because broad product-use categories overlap, but the intended score was not "
            "always the largest raw threshold-based response. Part of this result reflected unequal "
            "operating points: permissive thresholds generated high responses in both other positive sets and their own "
            "retained comparison sets. Subtracting the latter response improved diagonal dominance but did not eliminate "
            "cross-response. The functions should therefore not be described "
            "as mutually exclusive classifiers. A molecule may legitimately receive high values from several scores, "
            "and each value is better interpreted as the strength of evidence associated with that category."
        ),
        (
            "The component analysis supports inspectability without implying universal synergy. Combining descriptor "
            "and structural evidence improved AUC in most retained categories. "
            "Consequently, the component outputs can identify which evidence drives a score, but the combined form is "
            "not uniformly superior in every category."
        ),
        (
            "The bootstrap intervals show that the internal AUC advantage over QED was not attributable to a small number "
            "of sampled structures. They do not, however, remove dependence on category definitions or comparison-set "
            "construction. The external distributions reinforce this distinction: transfer was weak for endocrine disruptors "
            "and uncertain for pesticides because few nonoverlapping resolved structures "
            "were available."
        ),
        (
            "Exact-overlap-controlled revalidation prevented industrial co-use from being used as a post hoc explanation "
            "for nonspecific scores. Surfactants are legitimate cosmetic ingredients and can be authorized for food-contact "
            "uses, but these relationships do not validate a score that responds more often to category-exclusive "
            "surfactants than to its own positives. Official sources do confirm surfactant functions for cosmetic ingredients "
            "and surfactant or emulsifier uses in food-contact materials [28,29], but industrial plausibility does not establish "
            "analytical specificity. Cosmetics, food-contact substances, food additives, and solvents were therefore "
            "excluded. Animal-drug and human-drug functions were removed because QED already covers generic drug-likeness. "
            "Flavoring agents and fragrances were merged because their positive sets and scoring logic substantially overlapped; "
            "the merged target then passed three held-out evaluations."
        ),
    ]


def hard_negative_rebuild_text() -> list[str]:
    promotion_path = ROOT / "results" / "sequential_hard_negative_rebuild" / "cosmetics" / "promotion_report.json"
    if not promotion_path.is_file():
        return []
    promotion = json.loads(promotion_path.read_text(encoding="utf-8"))
    return [
        (
            "The cross-category analysis prompted sequential hard-negative rebuilds for animal drugs, human drugs, "
            "food-contact substances, cosmetics, and food additives. Candidate negatives were drawn directly from the "
            "other category sets after canonical target-overlap removal, without removing near-positive analogues. "
            "Cosmetics initially improved consistently across three held-out hash folds: hard-cross-category AUC increased "
            f"by {float(promotion['hard_auc_delta_mean']):.3f} on average, and the mean deployed false-positive response "
            f"decreased from {100 * float(promotion['hard_deployed_fpr_baseline_mean']):.1f}% to "
            f"{100 * float(promotion['hard_deployed_fpr_candidate_mean']):.1f}%. This gain involved a sensitivity "
            f"tradeoff ({100 * float(promotion['hard_deployed_tpr_baseline_mean']):.1f}% to "
            f"{100 * float(promotion['hard_deployed_tpr_candidate_mean']):.1f}%). A stricter source-specific revalidation "
            "then showed that the promoted function still called 81.0% of surfactants not assigned to cosmetics positive. "
            "It was therefore excluded together with the food-contact and food-additive functions."
        )
    ]


def replace_section(text: str, start: str, end: str, replacement: str) -> str:
    left, remainder = text.split(start, 1)
    _, right = remainder.split(end, 1)
    return left + replacement.rstrip() + "\n\n" + end + right


def build_markdown(
    base_path: Path,
    output_path: Path,
    rows: list[dict[str, str]],
    overlap_rows: list[dict[str, str]],
    overlap_figure: Path,
    cross_rows: list[dict[str, str]],
    ablation_rows: list[dict[str, str]],
    bootstrap_rows: list[dict[str, str]],
    external_distribution_rows: list[dict[str, str]],
    figure_dir: Path,
) -> None:
    text = base_path.read_text(encoding="utf-8")
    text = text.replace(
        "After the maximum-similarity screen at Tanimoto below 0.3, retained negative counts ranged from 3,422 for human drugs to 11,815 for surfactants.",
        "After the maximum-similarity screen at Tanimoto below 0.3, retained negative counts ranged from 3,455 for endocrine disruptors to 11,815 for surfactants.",
    )
    text = text.replace(
        "Surfactants, solvents, flavoring agents, fragrances, and endocrine disruptors formed the strongest group, each with AUC above 0.92 except surfactants, which reached 0.9778.",
        "Surfactants and endocrine disruptors showed the strongest retained separation, while the merged flavor-and-fragrance function also substantially exceeded QED.",
    )
    text = text.replace(
        "Structural patterns remained active across the full panel, with pattern counts ranging from one for solvents to eleven for endocrine disruptors (Figure 3). Descriptor support was likewise variable, ranging from one selected descriptor in food-contact substances to eight descriptors in the promoted cosmetics function.",
        "Structural patterns remained active across the retained panel, with category-specific motif sets summarized in Figure 3. Descriptor support likewise varied among the four retained functions.",
    )
    text = text.replace(
        "Flavoring agents and food additives share aldehyde, cinnamate, and ester patterns. Fragrances share aldehyde, cinnamate, ester, and long-chain patterns but add benzophenone-like support.",
        "The merged flavor-and-fragrance function retained structural evidence characteristic of both original source sets.",
    )
    text = refresh_markdown_core_tables(text, ablation_rows, bootstrap_rows)
    abstract = f"## Abstract\n\n{abstract_text(rows)}\n\n"
    text = replace_section(text, "## Abstract", "**Keywords:**", abstract)

    old_objective = next(line for line in text.splitlines() if line.startswith("The objective of this study was to test whether"))
    new_objective = old_objective.replace(
        "and Figure 4 compares structural-pattern usage across the final panel.",
        "Figure 4 compares structural-pattern usage across the final panel, Table 3 reports comparison with external positive category assignments after structure-level overlap control, Figure 5 shows direct parent-structure overlap between the PubChem-derived category sets and two external sources, Figures 6 and 7 characterize cross-category response and component contributions, Table 4 reports bootstrap confidence intervals, and Figure 8 compares internal and nonoverlapping external score distributions.",
    )
    text = text.replace(old_objective, new_objective)

    method_a, method_b = methods_text(rows)
    external_methods = (
        "### External positive-set comparison and overlap control\n\n"
        + method_a
        + "\n\n"
        + method_b
        + "\n\n"
    )
    additional_methods = (
        "### Additional characterization of the frozen scores\n\n"
        + analysis_methods_text()
        + "\n\n"
    )
    text = text.replace("## Results\n", external_methods + additional_methods + "## Results\n", 1)

    result_a, result_b = result_paragraphs(rows)
    external_results = (
        "### Comparison with external database category assignments\n\n"
        + result_a
        + "\n\n"
        + result_b
        + "\n\n"
        + markdown_table(rows)
        + "\n\n"
        + overlap_evidence_text(overlap_rows)
        + "\n\n"
        + f"![Figure 5. Direct parent-structure overlap between PubChem-derived category sets and external positive sources.](figures/{overlap_figure.name})"
        + "\n\n"
    )
    additional_results = (
        "### Cross-category response, component contribution, and uncertainty\n\n"
        + sequential_rebuild_results()
        + "\n\n"
        + cross_category_results(cross_rows)
        + "\n\n"
        + f"![Figure 6. Threshold response of frozen scores across chemical categories.](figures/figure6_cross_category_score_matrix.png)"
        + "\n\n"
        + ablation_results(ablation_rows)
        + "\n\n"
        + f"![Figure 7. Contribution of descriptor and structural-evidence components.](figures/figure7_component_ablation.png)"
        + "\n\n"
        + bootstrap_results(bootstrap_rows)
        + "\n\n"
        + confidence_interval_markdown_table(bootstrap_rows)
        + "\n\n"
        + external_distribution_results(external_distribution_rows, rows)
        + "\n\n"
        + f"![Figure 8. Score distributions for nonoverlapping external positive sets.](figures/figure8_external_score_distributions.png)"
        + "\n\n"
    )
    text = text.replace("## Discussion\n", external_results + additional_results + "## Discussion\n", 1)

    discussion = "## Discussion\n\n" + "\n\n".join(
        discussion_text(rows, overlap_rows) + additional_discussion_text() + hard_negative_rebuild_text()
    )
    text = replace_section(text, "## Discussion", "## Conclusions", discussion)
    conclusions = "## Conclusions\n\n" + "\n\n".join(conclusions_text(rows))
    text = replace_section(text, "## Conclusions", "## Figure legends", conclusions)
    figure_5_legend = (
        "**Figure 5.** Direct parent-structure overlap between the PubChem-derived category sets and external positive "
        "sources. Counts are unique standardized parent structures after structure resolution and parent normalization. "
        "Percentages are the fractions of each external set represented in the corresponding PubChem-derived category set. "
        "The PMRA panel includes only records for which structures were resolved. Circle areas are schematic and are not "
        "proportional to set size."
    )
    figure_6_legend = (
        "**Figure 6.** Threshold response of frozen scores across chemical categories. Rows denote complete positive "
        "category sets and columns denote the applied score. (A) Cells report the fraction at or above each score's frozen "
        "threshold. (B) Cells report the percentage-point difference between that fraction and the threshold-positive "
        "fraction in the score's retained comparison set; positive values indicate enrichment above the corresponding "
        "comparison response. Cyan outlines identify the intended category-score pairs. Off-diagonal responses show "
        "cross-response among nonexclusive product-use categories."
    )
    figure_7_legend = (
        "**Figure 7.** Contribution of descriptor and structural-evidence components. ROC AUC values were calculated "
        "against the same reconstructed target and retained comparison sets. The endocrine structural-evidence component "
        "includes scaffold, fingerprint-similarity, and SMARTS terms. The dashed line marks random performance."
    )
    figure_8_legend = (
        "**Figure 8.** Score distributions for the PubChem-derived category sets, retained comparison structures, and "
        "nonoverlapping external positive sets. Horizontal black lines mark medians and red dashed lines mark frozen "
        "thresholds. Group sizes are shown below each distribution."
    )
    text = text.replace(
        "**Figure 4.** Cross-category comparison of retained structural patterns. Filled cells mark named SMARTS patterns retained in the corresponding final scorer.",
        "**Figure 4.** Cross-category comparison of retained structural patterns. Filled cells mark named SMARTS patterns retained in the corresponding final scorer.\n\n"
        + figure_5_legend
        + "\n\n"
        + figure_6_legend
        + "\n\n"
        + figure_7_legend
        + "\n\n"
        + figure_8_legend,
    )
    supporting_text = (
        "## Code Availability\n\n"
        "The desktop application, Python library, four released scoring-function definitions, and usage instructions "
        "are available from the Chemical Category Scorer GitHub repository "
        "(https://github.com/phdgil/chemical-category-scorer), release version 2.0.0. The released model panel is "
        "identical to the four scoring functions reported here: endocrine disruptors, flavor and fragrance, pesticides, "
        "and surfactants.\n\n"
        "## Supporting Information\n\n"
        "Supporting Information is provided as `supporting_information_overlap_analysis.docx`: Table S1 and Figure S1 "
        "report screening and disposition of all eleven attempted scoring functions; Table S2 and Figure S2 report "
        "pairwise exact-structure counts, directional category coverage, and Jaccard similarity; Table S3 and Figure S3 "
        "report the multiplicity of original category assignments across unique structures.\n\n"
    )
    if "## Supporting Information" not in text:
        text = text.replace("## References\n", supporting_text + "## References\n", 1)

    for reference in REFERENCE_ADDITIONS:
        if reference not in text:
            text = text.rstrip() + "\n" + reference + "\n"
    output_path.write_text(apply_terminology(text), encoding="utf-8")


def find_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def paragraph_index(document: Document, target) -> int:
    return next(index for index, paragraph in enumerate(document.paragraphs) if paragraph._p is target._p)


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_results_table(document: Document, anchor, rows: list[dict[str, str]]) -> None:
    headers = [
        "Category",
        "External source",
        "Initial",
        "Resolved",
        "Overlap excluded",
        "Duplicates",
        "Scored positives",
        "Recovered",
        "Recovery",
        "Evidence scope",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = document.tables[0].style
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        if row.get("category") not in DISPLAY_NAMES:
            continue
        cells = table.add_row().cells
        values = [
            DISPLAY_NAMES[row["category"]],
            row["external_source"],
            str(integer(row, "raw_candidates")),
            str(integer(row, "resolved_structures")),
            str(integer(row, "overlap_excluded")),
            str(integer(row, "duplicate_external_excluded")),
            str(integer(row, "scored_true_external")),
            str(integer(row, "recovered_above_threshold")),
            percent(row),
            row["evidence_scope"],
        ]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
    anchor._p.addprevious(table._tbl)


def add_confidence_interval_table(document: Document, anchor, rows: list[dict[str, str]]) -> None:
    metrics: dict[str, dict[str, dict[str, str]]] = {}
    for row in rows:
        metrics.setdefault(row["category"], {})[row["metric"]] = row
    headers = ["Category", "Score AUC (95% CI)", "QED AUC (95% CI)", "ΔAUC (95% CI)"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = document.tables[0].style
    table.autofit = True
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for category in DISPLAY_NAMES:
        cells = table.add_row().cells
        values = [DISPLAY_NAMES[category]]
        for metric in ("auc", "qed_auc", "auc_delta_vs_qed"):
            row = metrics[category][metric]
            values.append(
                f"{float(row['estimate']):.3f} "
                f"({float(row['ci_lower_95']):.3f}–{float(row['ci_upper_95']):.3f})"
            )
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
    anchor._p.addprevious(table._tbl)


def add_figure_before(anchor, path: Path, caption_text: str) -> None:
    figure_paragraph = add_paragraph_before(anchor, "")
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(path), width=Inches(6.8))
    caption = add_paragraph_before(anchor, caption_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph_before(anchor, text: str, style: str = "Normal"):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    if style == "Heading 2":
        for run in paragraph.runs:
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
    return paragraph


def build_docx(
    base_path: Path,
    output_path: Path,
    rows: list[dict[str, str]],
    overlap_rows: list[dict[str, str]],
    overlap_figure: Path,
    cross_rows: list[dict[str, str]],
    ablation_rows: list[dict[str, str]],
    bootstrap_rows: list[dict[str, str]],
    external_distribution_rows: list[dict[str, str]],
    figure_dir: Path,
) -> None:
    shutil.copyfile(base_path, output_path)
    document = Document(output_path)
    for paragraph in document.paragraphs:
        if "retained negative counts ranged from 3,422 for human drugs to 11,815 for surfactants" in paragraph.text:
            paragraph.text = paragraph.text.replace(
                "retained negative counts ranged from 3,422 for human drugs to 11,815 for surfactants",
                "retained negative counts ranged from 3,455 for endocrine disruptors to 11,815 for surfactants",
            )
        paragraph.text = paragraph.text.replace(
            "Surfactants, solvents, flavoring agents, fragrances, and endocrine disruptors formed the strongest group, each with AUC above 0.92 except surfactants, which reached 0.9778.",
            "Surfactants and endocrine disruptors showed the strongest retained separation, while the merged flavor-and-fragrance function also substantially exceeded QED.",
        )
        paragraph.text = paragraph.text.replace(
            "Structural patterns remained active across the full panel, with pattern counts ranging from one for solvents to eleven for endocrine disruptors (Figure 3). Descriptor support was likewise variable, ranging from one selected descriptor in food-contact substances to eight descriptors in the promoted cosmetics function.",
            "Structural patterns remained active across the retained panel, with category-specific motif sets summarized in Figure 3. Descriptor support likewise varied among the four retained functions.",
        )
        paragraph.text = paragraph.text.replace(
            "Flavoring agents and food additives share aldehyde, cinnamate, and ester patterns. Fragrances share aldehyde, cinnamate, ester, and long-chain patterns but add benzophenone-like support.",
            "The merged flavor-and-fragrance function retained structural evidence characteristic of both original source sets.",
        )
    refresh_docx_core_tables(document, ablation_rows, bootstrap_rows)

    replace_paragraph(find_paragraph(document, "Accessible continuous chemical category scores"), abstract_text(rows))
    objective = find_paragraph(document, "The objective of this study was to test whether")
    replace_paragraph(
        objective,
        objective.text.replace(
            "and Figure 4 compares structural-pattern usage across the final panel.",
            "Figure 4 compares structural-pattern usage across the final panel, Table 3 reports comparison with external positive category assignments after structure-level overlap control, Figure 5 shows direct parent-structure overlap between the PubChem-derived category sets and two external sources, Figures 6 and 7 characterize cross-category response and component contributions, Table 4 reports bootstrap confidence intervals, and Figure 8 compares internal and nonoverlapping external score distributions.",
        ),
    )

    results_heading = find_paragraph(document, "3. Results")
    method_a, method_b = methods_text(rows)
    add_paragraph_before(results_heading, "2.6. External Positive-Set Comparison and Overlap Control", "Heading 2")
    add_paragraph_before(results_heading, method_a)
    add_paragraph_before(results_heading, method_b)
    add_paragraph_before(results_heading, "2.7. Additional Characterization of the Frozen Scores", "Heading 2")
    add_paragraph_before(results_heading, analysis_methods_text())

    discussion_heading = find_paragraph(document, "4. Discussion")
    result_a, result_b = result_paragraphs(rows)
    add_paragraph_before(discussion_heading, "3.5. Comparison with External Database Category Assignments", "Heading 2")
    add_paragraph_before(discussion_heading, result_a)
    add_paragraph_before(discussion_heading, result_b)
    add_paragraph_before(discussion_heading, overlap_evidence_text(overlap_rows))
    add_paragraph_before(
        discussion_heading,
        "Table 3. Comparison with external positive category assignments after structure resolution and overlap control. Recovery is the fraction of scored external positives at or above the pre-established category threshold.",
    )
    add_results_table(document, discussion_heading, rows)
    figure_paragraph = add_paragraph_before(discussion_heading, "")
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(overlap_figure), width=Inches(6.8))
    caption = add_paragraph_before(
        discussion_heading,
        "Figure 5. Direct parent-structure overlap between the PubChem-derived category sets and external positive sources. "
        "Counts are unique standardized parent structures after structure resolution and parent normalization. Percentages "
        "are the fractions of each external set represented in the corresponding PubChem-derived category set. The PMRA "
        "panel includes only records for which structures were resolved. Circle areas are schematic and are not proportional "
        "to set size.",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph_before(
        discussion_heading,
        "3.6. Cross-Category Response, Component Contribution, and Uncertainty",
        "Heading 2",
    )
    add_paragraph_before(discussion_heading, sequential_rebuild_results())
    add_paragraph_before(discussion_heading, cross_category_results(cross_rows))
    add_figure_before(
        discussion_heading,
        figure_dir / "figure6_cross_category_score_matrix.png",
        "Figure 6. Threshold response of frozen scores across chemical categories. Rows denote complete positive "
        "category sets and columns denote the applied score. (A) Cells report the fraction at or above each score's frozen "
        "threshold. (B) Cells report the percentage-point difference between that fraction and the threshold-positive "
        "fraction in the score's retained comparison set; positive values indicate enrichment above the corresponding "
        "comparison response. Cyan outlines identify the intended category-score pairs. Off-diagonal responses show "
        "cross-response among nonexclusive product-use categories.",
    )
    add_paragraph_before(discussion_heading, ablation_results(ablation_rows))
    add_figure_before(
        discussion_heading,
        figure_dir / "figure7_component_ablation.png",
        "Figure 7. Contribution of descriptor and structural-evidence components. ROC AUC values were calculated "
        "against the same reconstructed target and retained comparison sets. The endocrine structural-evidence component "
        "includes scaffold, fingerprint-similarity, and SMARTS terms. The dashed line marks random performance.",
    )
    add_paragraph_before(discussion_heading, bootstrap_results(bootstrap_rows))
    add_paragraph_before(
        discussion_heading,
        "Table 4. Bootstrap estimates and 95% confidence intervals from 1,000 stratified resamples of positive and "
        "retained comparison structures.",
    )
    add_confidence_interval_table(document, discussion_heading, bootstrap_rows)
    add_paragraph_before(
        discussion_heading,
        external_distribution_results(external_distribution_rows, rows),
    )
    add_figure_before(
        discussion_heading,
        figure_dir / "figure8_external_score_distributions.png",
        "Figure 8. Score distributions for the PubChem-derived category sets, retained comparison structures, and "
        "nonoverlapping external positive sets. Horizontal black lines mark medians and red dashed lines mark frozen "
        "thresholds. Group sizes are shown below each distribution.",
    )

    original_discussion = []
    start = paragraph_index(document, discussion_heading) + 1
    conclusion_heading = find_paragraph(document, "5. Conclusions")
    end = paragraph_index(document, conclusion_heading)
    original_discussion = document.paragraphs[start:end]
    replacements = discussion_text(rows, overlap_rows) + additional_discussion_text() + hard_negative_rebuild_text()
    for paragraph, replacement in zip(original_discussion, replacements):
        replace_paragraph(paragraph, replacement)
    if len(replacements) > len(original_discussion):
        for replacement in replacements[len(original_discussion):]:
            conclusion_heading.insert_paragraph_before(replacement, style="Normal")
    elif len(original_discussion) > len(replacements):
        for paragraph in original_discussion[len(replacements):]:
            paragraph._element.getparent().remove(paragraph._element)

    conclusion_heading = find_paragraph(document, "5. Conclusions")
    references_heading = find_paragraph(document, "References")
    conclusion_paragraphs = document.paragraphs[
        paragraph_index(document, conclusion_heading) + 1 : paragraph_index(document, references_heading)
    ]
    replacements = conclusions_text(rows)
    for paragraph, replacement in zip(conclusion_paragraphs, replacements):
        replace_paragraph(paragraph, replacement)
    for paragraph in conclusion_paragraphs[len(replacements):]:
        paragraph._element.getparent().remove(paragraph._element)

    references_heading = find_paragraph(document, "References")
    add_paragraph_before(references_heading, "Code Availability", "Heading 1")
    add_paragraph_before(
        references_heading,
        "The desktop application, Python library, four released scoring-function definitions, and usage instructions "
        "are available from the Chemical Category Scorer GitHub repository "
        "(https://github.com/phdgil/chemical-category-scorer), release version 2.0.0. The released model panel is "
        "identical to the four scoring functions reported here: endocrine disruptors, flavor and fragrance, pesticides, "
        "and surfactants.",
    )
    add_paragraph_before(references_heading, "Supporting Information", "Heading 1")
    add_paragraph_before(
        references_heading,
        "Supporting Information is provided as supporting_information_overlap_analysis.docx. Table S1 and Figure S1 "
        "report screening and disposition of all eleven attempted scoring functions; Table S2 and Figure S2 report "
        "pairwise exact-structure counts, directional category coverage, and Jaccard similarity; Table S3 and Figure S3 "
        "report the multiplicity of original category assignments across unique structures.",
    )

    first_reference = find_paragraph(document, "1. Bickerton")
    if "2. Ertl" in first_reference.text:
        replace_paragraph(first_reference, first_reference.text.split("2. Ertl", 1)[0].strip())

    for reference in REFERENCE_ADDITIONS:
        if not any(paragraph.text.strip() == reference for paragraph in document.paragraphs):
            document.add_paragraph(reference)

    for paragraph in document.paragraphs:
        revised = apply_terminology(paragraph.text)
        if revised != paragraph.text:
            replace_paragraph(paragraph, revised)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    revised = apply_terminology(paragraph.text)
                    if revised != paragraph.text:
                        replace_paragraph(paragraph, revised)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(12)
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.space_after = Pt(12)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.save(output_path)
    image_paths = [
        figure_dir / "figure1_final_rebuild_workflow.png",
        figure_dir / "figure2_auc_vs_qed.png",
        figure_dir / "figure3_descriptor_pattern_composition.png",
        figure_dir / "figure4_structural_pattern_comparison.png",
        overlap_figure,
        figure_dir / "figure6_cross_category_score_matrix.png",
        figure_dir / "figure7_component_ablation.png",
        figure_dir / "figure8_external_score_distributions.png",
    ]
    current_shapes = len(Document(output_path).inline_shapes)
    if current_shapes >= len(image_paths):
        replace_docx_inline_images(output_path, image_paths)
    elif current_shapes != 4:
        raise ValueError("DOCX has an unsupported number of pre-existing inline images")
    else:
        replace_docx_inline_images(output_path, image_paths[:4])
        completed = Document(output_path)
        discussion = find_paragraph(completed, "4. Discussion")
        for image_path in image_paths[4:]:
            paragraph = completed.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Inches(6.8))
            discussion._p.addprevious(paragraph._p)
        completed.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--summary", type=Path, required=True)
    parser.add_argument("--overlap-summary", type=Path, required=True)
    parser.add_argument("--overlap-figure", type=Path, required=True)
    parser.add_argument("--qed-analysis-dir", type=Path, required=True)
    parser.add_argument("--figure-dir", type=Path, required=True)
    parser.add_argument("--base-md", type=Path, required=True)
    parser.add_argument("--base-docx", type=Path, required=True)
    parser.add_argument("--output-md", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    rows = read_rows(args.summary)
    overlap_rows = read_rows(args.overlap_summary)
    cross_rows = read_rows(args.qed_analysis_dir / "cross_category_score_matrix.csv")
    ablation_rows = read_rows(args.qed_analysis_dir / "component_ablation.csv")
    bootstrap_rows = read_rows(args.qed_analysis_dir / "bootstrap_confidence_intervals.csv")
    external_distribution_rows = read_rows(args.qed_analysis_dir / "external_score_distributions.csv")
    if not primary_rows(rows):
        raise ValueError("No external positive comparison rows with scored structures were found")
    if not overlap_rows:
        raise ValueError("No direct source-overlap rows were found")
    expected_categories = len(DISPLAY_NAMES)
    if (
        len(cross_rows) != expected_categories**2
        or len(ablation_rows) != expected_categories * 3
        or len(bootstrap_rows) != expected_categories * 5
    ):
        raise ValueError("QED-inspired analysis tables are incomplete")
    if not external_distribution_rows:
        raise ValueError("No external score distributions were found")
    build_markdown(
        args.base_md,
        args.output_md,
        rows,
        overlap_rows,
        args.overlap_figure,
        cross_rows,
        ablation_rows,
        bootstrap_rows,
        external_distribution_rows,
        args.figure_dir,
    )
    build_docx(
        args.base_docx,
        args.output_docx,
        rows,
        overlap_rows,
        args.overlap_figure,
        cross_rows,
        ablation_rows,
        bootstrap_rows,
        external_distribution_rows,
        args.figure_dir,
    )
    print(args.output_md)
    print(args.output_docx)


if __name__ == "__main__":
    main()
