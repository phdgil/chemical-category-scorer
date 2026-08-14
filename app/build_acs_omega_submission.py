from __future__ import annotations

import argparse
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def build_toc_graphic(png_path: Path, tif_path: Path) -> None:
    png_path.parent.mkdir(parents=True, exist_ok=True)
    plt.rcParams.update({"font.family": "Arial", "font.size": 7})
    figure, axis = plt.subplots(figsize=(3.25, 1.75), dpi=300)
    figure.patch.set_facecolor("white")
    axis.set_xlim(0, 1)
    axis.set_ylim(0, 1)
    axis.axis("off")

    navy = "#17324D"
    blue = "#4477AA"
    pale_blue = "#EAF1F7"
    orange = "#EE7733"
    pale_orange = "#FCEDE5"
    green = "#228833"
    gray = "#667580"
    line = "#C8D2D9"

    axis.text(
        0.5,
        0.955,
        "Explicit scores for broad chemical product classes",
        ha="center",
        va="top",
        fontsize=8.2,
        fontweight="bold",
        color=navy,
    )
    axis.plot([0.035, 0.965], [0.855, 0.855], color=line, linewidth=0.7)

    axis.add_patch(
        plt.Rectangle((0.035, 0.255), 0.225, 0.50, facecolor=pale_blue, edgecolor=blue, linewidth=1.0)
    )
    axis.text(0.1475, 0.705, "PubChem category sets", ha="center", va="center", fontsize=6.2, fontweight="bold", color=navy)
    dot_positions = [
        (0.075, 0.605), (0.115, 0.625), (0.155, 0.605), (0.195, 0.625),
        (0.095, 0.535), (0.135, 0.555), (0.175, 0.535), (0.215, 0.555),
        (0.085, 0.455), (0.135, 0.465), (0.185, 0.455),
    ]
    axis.scatter(
        [x for x, _ in dot_positions],
        [y for _, y in dot_positions],
        s=25,
        color=blue,
        edgecolor="white",
        linewidth=0.45,
        zorder=3,
    )
    axis.text(0.1475, 0.335, "11 product classes", ha="center", va="center", fontsize=7.0, fontweight="bold", color=navy)

    axis.annotate(
        "",
        xy=(0.325, 0.505),
        xytext=(0.272, 0.505),
        arrowprops={"arrowstyle": "-|>", "color": gray, "lw": 1.1},
    )

    axis.add_patch(
        plt.Rectangle((0.335, 0.475), 0.255, 0.155, facecolor=pale_blue, edgecolor=blue, linewidth=1.0)
    )
    axis.text(0.4625, 0.5525, "Descriptor ranges", ha="center", va="center", fontsize=6.6, fontweight="bold", color=navy)
    axis.add_patch(
        plt.Rectangle((0.335, 0.275), 0.255, 0.155, facecolor=pale_orange, edgecolor=orange, linewidth=1.0)
    )
    axis.text(0.4625, 0.3525, "Structural patterns", ha="center", va="center", fontsize=6.6, fontweight="bold", color=navy)
    axis.text(0.4625, 0.4525, "+", ha="center", va="center", fontsize=9, fontweight="bold", color=gray)
    axis.text(0.4625, 0.70, "Inspectable score", ha="center", va="center", fontsize=7.0, fontweight="bold", color=navy)

    axis.annotate(
        "",
        xy=(0.657, 0.505),
        xytext=(0.602, 0.505),
        arrowprops={"arrowstyle": "-|>", "color": gray, "lw": 1.1},
    )

    axis.text(0.815, 0.73, "External-set recovery", ha="center", va="center", fontsize=6.8, fontweight="bold", color=navy)
    categories = ["Endocrine", "Drug", "Pesticide"]
    recoveries = [0.297, 0.876, 0.200]
    colors = [orange, green, orange]
    y_positions = [0.59, 0.46, 0.33]
    for label, recovery, color, y_position in zip(categories, recoveries, colors, y_positions):
        axis.text(0.665, y_position, label, ha="left", va="center", fontsize=5.6, color=navy)
        axis.add_patch(
            plt.Rectangle((0.755, y_position - 0.035), 0.18, 0.07, facecolor="#EDF1F3", edgecolor="none")
        )
        axis.add_patch(
            plt.Rectangle((0.755, y_position - 0.035), 0.18 * recovery, 0.07, facecolor=color, edgecolor="none")
        )
        axis.text(0.947, y_position, f"{100 * recovery:.0f}%", ha="right", va="center", fontsize=5.7, fontweight="bold", color=navy)
    axis.text(0.815, 0.235, "Transfer varies by category", ha="center", va="center", fontsize=6.5, fontweight="bold", color=navy)

    figure.subplots_adjust(left=0.005, right=0.995, bottom=0.02, top=0.985)
    figure.savefig(png_path, dpi=300, facecolor="white")
    figure.savefig(tif_path, dpi=300, facecolor="white", pil_kwargs={"compression": "tiff_lzw"})
    plt.close(figure)


def prepare_imagegen_graphic(source: Path, png_path: Path, tif_path: Path) -> None:
    png_path.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(source) as image:
        image = image.convert("RGB")
        if image.size != (1950, 1050):
            image = image.resize((1950, 1050), Image.Resampling.LANCZOS)
        image.save(png_path, "PNG", dpi=(600, 600), optimize=True)
        image.save(tif_path, "TIFF", dpi=(600, 600), compression="tiff_lzw")


def build_docx(source: Path, output: Path, toc_png: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, output)
    document = Document(output)
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    label = document.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("For Table of Contents Only")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    graphic = document.add_paragraph()
    graphic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    graphic.add_run().add_picture(str(toc_png), width=Inches(3.25))
    document.save(output)


def build_markdown(source: Path, output: Path, toc_png: Path) -> None:
    text = source.read_text(encoding="utf-8").rstrip()
    text += f"\n\n## For Table of Contents Only\n\n![ACS Omega TOC graphic](figures/{toc_png.name})\n"
    output.write_text(text, encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-docx", type=Path, required=True)
    parser.add_argument("--source-md", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-md", type=Path, required=True)
    parser.add_argument("--toc-png", type=Path, required=True)
    parser.add_argument("--toc-tif", type=Path, required=True)
    parser.add_argument("--toc-source", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if args.toc_source:
        prepare_imagegen_graphic(args.toc_source, args.toc_png, args.toc_tif)
    else:
        build_toc_graphic(args.toc_png, args.toc_tif)
    build_docx(args.source_docx, args.output_docx, args.toc_png)
    build_markdown(args.source_md, args.output_md, args.toc_png)
    print(args.output_docx)
    print(args.output_md)
    print(args.toc_png)
    print(args.toc_tif)


if __name__ == "__main__":
    main()
