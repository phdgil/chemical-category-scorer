from __future__ import annotations

import csv
import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from rdkit import Chem

from app.qed_inspired_validation import ROOT, category_positive_smiles

ANALYSIS = ROOT / "results" / "qed_inspired_analysis"
PAPER = ROOT / "paper"

DISPLAY = {
    "animal_drugs": "Animal drugs",
    "human_drugs": "Human drugs",
    "cosmetics": "Cosmetics",
    "food_contact_substances": "Food contact",
    "food_additives": "Food additives",
    "solvents": "Solvents",
    "flavoring_agents": "Flavoring agents",
    "fragrances": "Fragrances",
    "endocrine_disruptors": "Endocrine disruptors",
    "pesticides": "Pesticides",
    "surfactants": "Surfactants",
}
CATEGORIES = tuple(DISPLAY)
FINAL_DISPOSITION = {
    "animal_drugs": "exclude",
    "human_drugs": "exclude",
    "cosmetics": "exclude",
    "endocrine_disruptors": "retain",
    "flavoring_agents": "merge as flavor and fragrance",
    "food_additives": "exclude",
    "food_contact_substances": "exclude",
    "fragrances": "merge as flavor and fragrance",
    "pesticides": "retain",
    "solvents": "exclude",
    "surfactants": "retain",
}
DISPOSITION_REASON = {
    "animal_drugs": "Broad drug-use class removed; QED already represents generic drug-likeness and hard-negative rebuilding did not improve discrimination.",
    "human_drugs": "Broad drug-use class removed; QED already represents generic drug-likeness and combined-source rebuilding damaged the original benchmark.",
    "cosmetics": "Worst exact-overlap-excluded source response (surfactants, 81.0%) exceeded the own-category response (43.5%).",
    "endocrine_disruptors": "Retained: hard-background AUC 0.917 and worst foreign response 32.0%.",
    "flavoring_agents": "Merged with fragrances because 46.2% of flavoring structures occurred in both sets and reciprocal exclusive-set responses were high.",
    "food_additives": "Worst exact-overlap-excluded source response (flavoring agents, 90.2%) exceeded the own-category response (76.7%).",
    "food_contact_substances": "Hard-background AUC was 0.508 and the surfactant response reached 100%.",
    "fragrances": "Merged with flavoring agents because 46.9% of fragrance structures occurred in both sets and reciprocal exclusive-set responses were high.",
    "pesticides": "Retained after network augmentation: hard-background AUC 0.741 and worst foreign response 34.5%.",
    "solvents": "Excluded because 71.8% of exact-overlap-excluded flavoring agents met the solvent threshold, above the 50% final specificity ceiling.",
    "surfactants": "Retained: hard-background AUC 0.964 and worst foreign response 21.6%.",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def canonical_positive_sets() -> dict[str, set[str]]:
    output: dict[str, set[str]] = {}
    for category in CATEGORIES:
        values: set[str] = set()
        for smiles in category_positive_smiles(category):
            molecule = Chem.MolFromSmiles(smiles)
            if molecule is not None:
                values.add(Chem.MolToSmiles(molecule, canonical=True))
        output[category] = values
    return output


def build_overlap_assets() -> tuple[Path, Path, Path, Path]:
    sets = canonical_positive_sets()
    pairwise_rows: list[dict[str, object]] = []
    directional = np.zeros((len(CATEGORIES), len(CATEGORIES)), dtype=float)
    for row_index, category_a in enumerate(CATEGORIES):
        for column_index, category_b in enumerate(CATEGORIES):
            shared = len(sets[category_a] & sets[category_b])
            directional[row_index, column_index] = shared / len(sets[category_a])
            if row_index >= column_index:
                continue
            union = len(sets[category_a] | sets[category_b])
            pairwise_rows.append(
                {
                    "category_a": category_a,
                    "category_b": category_b,
                    "category_a_count": len(sets[category_a]),
                    "category_b_count": len(sets[category_b]),
                    "shared_structure_count": shared,
                    "fraction_of_category_a": shared / len(sets[category_a]),
                    "fraction_of_category_b": shared / len(sets[category_b]),
                    "jaccard_similarity": shared / union,
                }
            )
    pairwise_path = PAPER / "supporting_information_pairwise_category_overlap.csv"
    with pairwise_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(pairwise_rows[0]))
        writer.writeheader()
        writer.writerows(pairwise_rows)

    labels = [DISPLAY[category] for category in CATEGORIES]
    fig, ax = plt.subplots(figsize=(14, 11), constrained_layout=True)
    image = ax.imshow(directional, cmap="YlOrRd", vmin=0, vmax=1)
    for row_index in range(len(CATEGORIES)):
        for column_index in range(len(CATEGORIES)):
            value = directional[row_index, column_index]
            ax.text(
                column_index,
                row_index,
                f"{100 * value:.0f}%",
                ha="center",
                va="center",
                fontsize=9,
                color="white" if value >= 0.55 else "black",
            )
    ax.set_xticks(range(len(CATEGORIES)), labels, rotation=45, ha="right")
    ax.set_yticks(range(len(CATEGORIES)), labels)
    ax.set_xlabel("Comparison category")
    ax.set_ylabel("Denominator category")
    ax.set_title("Directional exact-structure overlap among the eleven original positive sets", fontsize=19, weight="bold")
    colorbar = fig.colorbar(image, ax=ax, fraction=0.04, pad=0.03)
    colorbar.set_label("Fraction of row-category structures also assigned to column category")
    overlap_figure = PAPER / "figures" / "figureS2_directional_category_overlap.png"
    fig.savefig(overlap_figure, dpi=300, facecolor="white")
    plt.close(fig)

    all_structures = set().union(*sets.values())
    multiplicity_counts: dict[int, int] = {}
    for smiles in all_structures:
        multiplicity = sum(smiles in values for values in sets.values())
        multiplicity_counts[multiplicity] = multiplicity_counts.get(multiplicity, 0) + 1
    multiplicity_rows = [
        {
            "category_assignment_count": assignment_count,
            "unique_structure_count": structure_count,
            "fraction_of_all_unique_structures": structure_count / len(all_structures),
        }
        for assignment_count, structure_count in sorted(multiplicity_counts.items())
    ]
    multiplicity_path = PAPER / "supporting_information_category_multiplicity.csv"
    with multiplicity_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(multiplicity_rows[0]))
        writer.writeheader()
        writer.writerows(multiplicity_rows)

    fig, ax = plt.subplots(figsize=(10, 7), constrained_layout=True)
    bars = ax.bar(
        [row["category_assignment_count"] for row in multiplicity_rows],
        [row["unique_structure_count"] for row in multiplicity_rows],
        color="#4C78A8",
        edgecolor="#333333",
    )
    for bar, row in zip(bars, multiplicity_rows):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + max(multiplicity_counts.values()) * 0.012,
            f"{int(row['unique_structure_count']):,}",
            ha="center",
            va="bottom",
            fontsize=10,
        )
    ax.set_yscale("log")
    ax.set_xticks(sorted(multiplicity_counts))
    ax.set_xlabel("Number of original categories assigned to a structure")
    ax.set_ylabel("Unique structure count (log scale)")
    ax.set_title("Multiplicity of category assignments across unique structures", fontsize=19, weight="bold")
    ax.grid(axis="y", alpha=0.25)
    multiplicity_figure = PAPER / "figures" / "figureS3_category_assignment_multiplicity.png"
    fig.savefig(multiplicity_figure, dpi=300, facecolor="white")
    plt.close(fig)
    return pairwise_path, overlap_figure, multiplicity_path, multiplicity_figure


def add_csv_table(document: Document, path: Path, columns: list[tuple[str, str]]) -> None:
    rows = read_csv(path)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, (_key, label) in zip(table.rows[0].cells, columns):
        cell.text = label
    for row in rows:
        cells = table.add_row().cells
        for cell, (key, _label) in zip(cells, columns):
            value = row.get(key, "")
            if key.startswith("fraction") or key == "jaccard_similarity":
                value = f"{100 * float(value):.1f}%" if value else ""
            elif key.endswith("prevalence"):
                value = f"{100 * float(value):.1f}%" if value else ""
            elif key in {"own_threshold_response", "worst_source_threshold_response"}:
                value = f"{100 * float(value):.1f}%" if value else ""
            elif key == "exact_overlap_excluded_hard_auc":
                value = f"{float(value):.3f}" if value else ""
            elif "auc_delta" in key:
                value = f"{float(value):+.3f}" if value else ""
            elif key in {"mean_walk_enrichment", "mean_weight"}:
                value = f"{float(value):.3f}" if value else ""
            cell.text = value


def build_supporting_document(
    score_selection: Path,
    screening_figure: Path,
    pairwise_path: Path,
    overlap_figure: Path,
    multiplicity_path: Path,
    multiplicity_figure: Path,
    external_database_path: Path,
    combined_test_path: Path,
    combined_test_figure: Path,
    network_performance_path: Path,
    network_pattern_path: Path,
    network_figure: Path,
) -> Path:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    title = document.add_heading("Supporting Information", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Chemical Category Scoring across Broad Use and Hazard Classes Using Molecular Descriptors and Structural Patterns",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("Candidate-score screening", level=1)
    document.add_paragraph(
        "Scoring-function development was attempted for all eleven original categories. Table S1 reports the "
        "exact-overlap-controlled screening result and final disposition of every candidate. Standalone functions were "
        "retained only when hard-background AUC exceeded 0.5, the own-category response exceeded every foreign-category "
        "response, and no exact-overlap-excluded foreign category reached 50%. Strongly reciprocal flavoring-agent and "
        "fragrance responses were handled by evaluating a merged target rather than retaining two nonspecific functions."
    )
    document.add_paragraph(
        "Table S1. Publication-screening results for all attempted category scoring functions."
    )
    add_csv_table(
        document,
        score_selection,
        [
            ("category", "Category"),
            ("own_threshold_response", "Own response"),
            ("worst_exact_overlap_excluded_source", "Worst foreign source"),
            ("worst_source_threshold_response", "Worst response"),
            ("exact_overlap_excluded_hard_auc", "Hard AUC"),
            ("final_disposition", "Disposition"),
            ("disposition_reason", "Reason"),
        ],
    )
    document.add_picture(str(screening_figure), width=Inches(6.7))
    document.add_paragraph(
        "Figure S1. Own-category threshold response and worst exact-overlap-excluded foreign-category response for "
        "quantitatively revalidated standalone candidates. Drug-category exclusions and the flavor-and-fragrance merger "
        "are reported in Table S1."
    )
    document.add_heading("Exact structural overlap among original categories", level=1)
    document.add_paragraph(
        "Canonical SMILES were deduplicated within each original positive set. For every category pair, Table S2 "
        "reports the exact shared-structure count, the fraction of each category represented by the intersection, "
        "and Jaccard similarity. Figure S2 displays directional overlap, with each cell using the row category as "
        "the denominator."
    )
    document.add_paragraph("Table S2. Pairwise exact-structure overlap among the eleven original positive sets.")
    add_csv_table(
        document,
        pairwise_path,
        [
            ("category_a", "Category A"),
            ("category_b", "Category B"),
            ("shared_structure_count", "Shared"),
            ("fraction_of_category_a", "% of A"),
            ("fraction_of_category_b", "% of B"),
            ("jaccard_similarity", "Jaccard"),
        ],
    )
    document.add_picture(str(overlap_figure), width=Inches(6.7))
    document.add_paragraph(
        "Figure S2. Directional exact-structure overlap among the eleven original positive sets. Values are the "
        "percentage of structures in the row category also assigned to the column category."
    )
    document.add_heading("Multiplicity of category assignments", level=1)
    document.add_paragraph(
        "Assignment multiplicity was counted over the union of unique canonical structures from all eleven positive sets."
    )
    document.add_paragraph("Table S3. Number of original category assignments per unique structure.")
    add_csv_table(
        document,
        multiplicity_path,
        [
            ("category_assignment_count", "Category assignments"),
            ("unique_structure_count", "Unique structures"),
            ("fraction_of_all_unique_structures", "Fraction of union"),
        ],
    )
    document.add_picture(str(multiplicity_figure), width=Inches(6.7))
    document.add_paragraph(
        "Figure S3. Multiplicity of original category assignments across unique structures. The vertical axis uses "
        "a logarithmic scale."
    )
    document.add_heading("External databases and combined-positive-set experiment", level=1)
    document.add_paragraph(
        "Table S4 lists every external database considered for category enrichment or external comparison. Resolved "
        "structures were standardized to unique parent structures when audited parent-level data were available and "
        "otherwise deduplicated as canonical structures. Target-PubChem overlap is distinct from usable additions: the "
        "latter also excludes duplicates and structures overlapping any construction set used by the corresponding "
        "experiment. Partial identifier resolution is reported explicitly."
    )
    document.add_paragraph("Table S4. External databases considered for positive-set enrichment.")
    add_csv_table(
        document,
        external_database_path,
        [
            ("category", "Category"),
            ("database", "External database"),
            ("citation", "Citation"),
            ("candidate_records", "Candidates"),
            ("resolved_records", "Resolved"),
            ("unique_resolved_parents", "Unique resolved structures"),
            ("target_pubchem_overlap", "Target-PubChem overlap"),
            ("usable_additions", "Usable additions"),
            ("use_in_final_test", "Use"),
        ],
    )
    document.add_paragraph(
        "The final experiment added external structures to the PubChem positives within each training fold and rebuilt "
        "the corresponding comparison background after exact target-overlap removal. External additions were divided "
        "into three molecular-hash folds. Table S5 reports candidate-minus-baseline AUC changes on held-out external "
        "positives versus held-out constructed hard cross-category negatives and on the original PubChem benchmark. "
        "These AUCs are not external-database specificity estimates because their negative class was constructed from "
        "other category sources. For endocrine disruption, the final experiment updated descriptor statistics and "
        "similarity references while retaining the frozen scaffold and SMARTS terms; the other candidates were rebuilt "
        "with the descriptor-range and structural-pattern pipeline."
    )
    document.add_paragraph("Table S5. Three-fold combined-source positive-set rebuilding results.")
    add_csv_table(
        document,
        combined_test_path,
        [
            ("category", "Category"),
            ("pubchem_positive_count", "PubChem positives"),
            ("external_new_count", "External additions"),
            ("mean_external_holdout_auc_delta", "Mean external ΔAUC"),
            ("minimum_external_holdout_auc_delta", "Minimum external ΔAUC"),
            ("mean_original_auc_delta", "Mean original ΔAUC"),
            ("promotion_gate", "Promotion gate"),
        ],
    )
    document.add_picture(str(combined_test_figure), width=Inches(6.7))
    document.add_paragraph(
        "Figure S4. Candidate-minus-baseline AUC changes after combined PubChem and external positive-set rebuilding. "
        "Bars show mean held-out external and original-benchmark changes across three molecular-hash folds."
    )
    document.add_heading("Network-analysis random-walk pattern discovery", level=1)
    document.add_paragraph(
        "A bipartite network was constructed independently in each molecular-hash training fold. Molecule nodes were "
        "connected to canonical atom-neighborhood fragment nodes generated at Morgan radii 1–3. Personalized PageRank "
        "with damping 0.85 was run separately from positive- and hard-negative-molecule restart distributions [36,37]. "
        "Fragments were ranked by the ratio of positive- to negative-seeded stationary probability, subject to minimum "
        "positive prevalence and prevalence-enrichment filters. The resulting network-pattern score was mixed with each "
        "frozen baseline using a training-fold weight; all reported AUC changes were measured in the held-out fold."
    )
    document.add_paragraph("Table S6. Three-fold network random-walk pattern-augmentation results.")
    add_csv_table(
        document,
        network_performance_path,
        [
            ("category", "Category"),
            ("mean_hard_auc_delta", "Mean hard ΔAUC"),
            ("minimum_hard_auc_delta", "Minimum hard ΔAUC"),
            ("mean_original_auc_delta", "Mean original ΔAUC"),
            ("promotion_gate", "Promotion gate"),
        ],
    )
    document.add_paragraph(
        "Only pesticides passed the promotion gate. Fifteen fragment neighborhoods appeared among the top 24 in all "
        "three folds and were retained in the full-data definition."
    )
    document.add_paragraph("Table S7. Consensus pesticide fragments discovered by network random walk.")
    add_csv_table(
        document,
        network_pattern_path,
        [
            ("fragment_smiles", "Fragment"),
            ("mean_positive_prevalence", "Mean positive prevalence"),
            ("mean_walk_enrichment", "Mean walk enrichment"),
            ("mean_weight", "Mean weight"),
        ],
    )
    document.add_picture(str(network_figure), width=Inches(6.7))
    document.add_paragraph(
        "Figure S5. Held-out AUC changes from bipartite-network random-walk pattern augmentation. Pesticides improved "
        "in every fold and passed the mean-improvement and original-benchmark preservation gates."
    )
    output = PAPER / "supporting_information_final.docx"
    document.save(output)
    return output


def build_external_database_assets() -> tuple[Path, Path, Path]:
    rows = [
        {"category": "endocrine_disruptors", "database": "DEDuCT v3 categories I–III", "citation": "[25]", "candidate_records": 704, "resolved_records": 704, "unique_resolved_parents": 673, "target_pubchem_overlap": 536, "usable_additions": 64, "use_in_final_test": "yes"},
        {"category": "flavor_fragrance", "database": "EU Union List of Flavouring Substances", "citation": "[29]", "candidate_records": 2505, "resolved_records": 2264, "unique_resolved_parents": 2250, "target_pubchem_overlap": 1843, "usable_additions": 407, "use_in_final_test": "yes"},
        {"category": "pesticides", "database": "Health Canada PMRA PPID", "citation": "[26]", "candidate_records": 545, "resolved_records": 228, "unique_resolved_parents": 204, "target_pubchem_overlap": 163, "usable_additions": 15, "use_in_final_test": "yes"},
        {"category": "surfactants", "database": "US EPA Safer Chemical Ingredients List", "citation": "[30]", "candidate_records": 355, "resolved_records": 181, "unique_resolved_parents": 174, "target_pubchem_overlap": 12, "usable_additions": 162, "use_in_final_test": "yes"},
        {"category": "animal_drugs", "database": "Health Canada Drug Product Database", "citation": "[33]", "candidate_records": 443, "resolved_records": 290, "unique_resolved_parents": 284, "target_pubchem_overlap": 101, "usable_additions": 183, "use_in_final_test": "yes"},
        {"category": "human_drugs", "database": "DrugCentral", "citation": "[32]", "candidate_records": 4099, "resolved_records": 4099, "unique_resolved_parents": 4013, "target_pubchem_overlap": 2001, "usable_additions": 1633, "use_in_final_test": "yes"},
        {"category": "human_drugs", "database": "Health Canada Drug Product Database", "citation": "[33]", "candidate_records": 1887, "resolved_records": 14, "unique_resolved_parents": 14, "target_pubchem_overlap": 7, "usable_additions": 6, "use_in_final_test": "partial cached resolution"},
        {"category": "food_additives", "database": "Health Canada Lists of Permitted Food Additives", "citation": "[31]", "candidate_records": 498, "resolved_records": 239, "unique_resolved_parents": 236, "target_pubchem_overlap": 151, "usable_additions": 85, "use_in_final_test": "yes"},
        {"category": "solvents", "database": "US EPA Safer Chemical Ingredients List", "citation": "[30]", "candidate_records": 117, "resolved_records": 101, "unique_resolved_parents": 100, "target_pubchem_overlap": 25, "usable_additions": 75, "use_in_final_test": "yes"},
        {"category": "cosmetics", "database": "California Safe Cosmetics Program", "citation": "[34]", "candidate_records": 113, "resolved_records": 4, "unique_resolved_parents": 4, "target_pubchem_overlap": 4, "usable_additions": 0, "use_in_final_test": "no new structures"},
        {"category": "food_contact_substances", "database": "FSA/FSS regulated-products register", "citation": "[35]", "candidate_records": 763, "resolved_records": 0, "unique_resolved_parents": 0, "target_pubchem_overlap": 0, "usable_additions": 0, "use_in_final_test": "unresolved"},
    ]
    database_path = PAPER / "supporting_information_external_databases.csv"
    with database_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)

    summary = json.loads(
        (ROOT / "results/combined_external_positive_rebuild/summary.json").read_text(encoding="utf-8")
    )
    test_rows: list[dict[str, object]] = []
    external_counts = {
        row["category"]: int(row["usable_additions"])
        for row in rows
        if row["use_in_final_test"] == "yes"
    }
    external_counts["human_drugs"] = 1639
    for category, values in summary.items():
        folds = values["folds"]
        test_rows.append(
            {
                "category": category,
                "pubchem_positive_count": folds[0]["pubchem_positive_count"],
                "external_new_count": external_counts[category],
                "mean_external_holdout_auc_delta": values["mean_external_holdout_auc_delta"],
                "minimum_external_holdout_auc_delta": values["minimum_external_holdout_auc_delta"],
                "mean_original_auc_delta": values["mean_original_auc_delta"],
                "promotion_gate": str(values["promotion_gate"]).lower(),
            }
        )
    test_path = PAPER / "supporting_information_combined_positive_rebuild.csv"
    with test_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(test_rows[0]))
        writer.writeheader()
        writer.writerows(test_rows)

    x = np.arange(len(test_rows))
    width = 0.36
    fig, ax = plt.subplots(figsize=(14, 8), constrained_layout=True)
    ax.bar(
        x - width / 2,
        [float(row["mean_external_holdout_auc_delta"]) for row in test_rows],
        width,
        label="Held-out external ΔAUC",
        color="#4C78A8",
    )
    ax.bar(
        x + width / 2,
        [float(row["mean_original_auc_delta"]) for row in test_rows],
        width,
        label="Original-benchmark ΔAUC",
        color="#F58518",
    )
    ax.axhline(0, color="black", linewidth=1)
    ax.set_xticks(x, [DISPLAY.get(row["category"], row["category"]).replace("_", " ") for row in test_rows], rotation=40, ha="right")
    ax.set_ylabel("Candidate minus baseline AUC")
    ax.set_title("Combined PubChem and external positive-set rebuilding", fontsize=20, weight="bold")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    figure_path = PAPER / "figures/figureS4_combined_positive_rebuild.png"
    fig.savefig(figure_path, dpi=300, facecolor="white")
    plt.close(fig)
    return database_path, test_path, figure_path


def build_network_random_walk_assets() -> tuple[Path, Path, Path]:
    summary = json.loads(
        (ROOT / "results/network_random_walk_pattern_rebuild/summary.json").read_text(encoding="utf-8")
    )
    rows = [
        {
            "category": category,
            "mean_hard_auc_delta": values["mean_hard_auc_delta"],
            "minimum_hard_auc_delta": values["minimum_hard_auc_delta"],
            "mean_original_auc_delta": values["mean_original_auc_delta"],
            "promotion_gate": str(values["promotion_gate"]).lower(),
        }
        for category, values in summary.items()
    ]
    performance_path = PAPER / "supporting_information_network_random_walk.csv"
    with performance_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)

    candidate = json.loads(
        (ROOT / "results/network_random_walk_pattern_rebuild/final_pesticide_candidate.json").read_text(
            encoding="utf-8"
        )
    )
    pattern_rows = [
        {
            "fragment_smiles": row["fragment_smiles"],
            "mean_positive_prevalence": row["mean_positive_prevalence"],
            "mean_walk_enrichment": row["mean_walk_enrichment"],
            "mean_weight": row["weight"],
        }
        for row in candidate["patterns"]
    ]
    pattern_path = PAPER / "supporting_information_pesticide_network_patterns.csv"
    with pattern_path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(pattern_rows[0]))
        writer.writeheader()
        writer.writerows(pattern_rows)

    x = np.arange(len(rows))
    width = 0.36
    fig, ax = plt.subplots(figsize=(11, 7), constrained_layout=True)
    ax.bar(
        x - width / 2,
        [float(row["mean_hard_auc_delta"]) for row in rows],
        width,
        label="Held-out hard-background ΔAUC",
        color="#4C78A8",
    )
    ax.bar(
        x + width / 2,
        [float(row["mean_original_auc_delta"]) for row in rows],
        width,
        label="Original-benchmark ΔAUC",
        color="#F58518",
    )
    ax.axhline(0, color="black", linewidth=1)
    ax.axhline(0.02, color="#555555", linestyle="--", linewidth=1, label="Mean promotion gate")
    ax.set_xticks(
        x,
        [DISPLAY.get(row["category"], row["category"]).replace("_", " ") for row in rows],
        rotation=30,
        ha="right",
    )
    ax.set_ylabel("Network-augmented minus baseline AUC")
    ax.set_title("Bipartite-network random-walk pattern augmentation", fontsize=19, weight="bold")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    figure_path = PAPER / "figures/figureS5_network_random_walk_pattern_augmentation.png"
    fig.savefig(figure_path, dpi=300, facecolor="white")
    plt.close(fig)
    return performance_path, pattern_path, figure_path


def main() -> None:
    rows = read_csv(ANALYSIS / "publication_specificity_revalidation.csv")
    animal_human = [
        {
            "category": "animal_drugs",
            "own_threshold_response": "",
            "worst_source_threshold_response": "",
            "exact_overlap_excluded_hard_auc": "",
            "publication_decision": "exclude: QED overlap",
        },
        {
            "category": "human_drugs",
            "own_threshold_response": "",
            "worst_source_threshold_response": "",
            "exact_overlap_excluded_hard_auc": "",
            "publication_decision": "exclude: QED overlap",
        },
    ]
    output_rows = animal_human + rows
    table_path = PAPER / "supporting_information_score_selection.csv"
    with table_path.open("w", encoding="utf-8-sig", newline="") as handle:
        fields = [
            "category",
            "own_threshold_response",
            "worst_exact_overlap_excluded_source",
            "worst_source_threshold_response",
            "exact_overlap_excluded_hard_auc",
            "publication_decision",
            "final_disposition",
            "disposition_reason",
        ]
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        for row in output_rows:
            category = row["category"]
            writer.writerow(
                {field: row.get(field, "") for field in fields[:-2]}
                | {
                    "final_disposition": FINAL_DISPOSITION[category],
                    "disposition_reason": DISPOSITION_REASON[category],
                }
            )

    plot_rows = [row for row in rows if row["category"] not in {"flavoring_agents", "fragrances"}]
    categories = [row["category"] for row in plot_rows]
    own = np.asarray([float(row["own_threshold_response"]) for row in plot_rows])
    worst = np.asarray([float(row["worst_source_threshold_response"]) for row in plot_rows])
    x = np.arange(len(categories)); width = 0.36
    fig, ax = plt.subplots(figsize=(14, 8), constrained_layout=True)
    ax.bar(x - width / 2, own, width, color="#2C7FB8", label="Own positive set")
    ax.bar(x + width / 2, worst, width, color="#F28E2B", label="Worst exact-overlap-excluded source")
    for index, row in enumerate(plot_rows):
        retained = FINAL_DISPOSITION[row["category"]] == "retain"
        ax.text(index, max(own[index], worst[index]) + 0.035, "Retain" if retained else "Exclude", ha="center", weight="bold", color="#1B7837" if retained else "#B2182B")
    ax.set_title("Specificity screening of standalone category-score candidates", fontsize=21, weight="bold")
    ax.set_ylabel("Fraction at or above frozen threshold", fontsize=15)
    ax.set_xticks(x, [DISPLAY[c] for c in categories], rotation=40, ha="right")
    ax.set_ylim(0, 1.12); ax.grid(axis="y", alpha=0.25); ax.legend(fontsize=12)
    fig.savefig(PAPER / "figures" / "figureS1_publication_score_screening.png", dpi=300, facecolor="white")
    plt.close(fig)
    pairwise_path, overlap_figure, multiplicity_path, multiplicity_figure = build_overlap_assets()
    external_database_path, combined_test_path, combined_test_figure = build_external_database_assets()
    network_performance_path, network_pattern_path, network_figure = build_network_random_walk_assets()
    supporting_document = build_supporting_document(
        table_path,
        PAPER / "figures" / "figureS1_publication_score_screening.png",
        pairwise_path,
        overlap_figure,
        multiplicity_path,
        multiplicity_figure,
        external_database_path,
        combined_test_path,
        combined_test_figure,
        network_performance_path,
        network_pattern_path,
        network_figure,
    )
    print(table_path)
    print(PAPER / "figures" / "figureS1_publication_score_screening.png")
    print(pairwise_path)
    print(overlap_figure)
    print(multiplicity_path)
    print(multiplicity_figure)
    print(supporting_document)
    print(external_database_path)
    print(combined_test_path)
    print(combined_test_figure)


if __name__ == "__main__":
    main()
